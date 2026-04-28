"""様式2-1 の素テンプレ (r1y2-1.docx) にプレースホルダを差し込み、
ランタイムが使う template.docx を生成するスクリプト。

使い方 (リポジトリルートから):
  uv run python backend/scripts/build_template.py

r1y2-1.docx を上書きしたい新版が来たときも、再実行すれば template.docx を作り直せる。
"""

from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[2]
SRC = REPO_ROOT / "backend" / "app" / "templates" / "r1y2-1.docx"
DST = REPO_ROOT / "backend" / "app" / "templates" / "template.docx"

# (テーブル index, row, col, プレースホルダ) の対応表。
# r1y2-1.docx のレイアウトに依存しているため、新版で表構造が変わったら要更新。
INSERTIONS = [
    # 申請者の概要: 「（フリガナ）/ 名称（商号または屋号）」セルの右側に会社名
    (0, 0, 1, "{{company}}"),
    # 経営計画: 「4. 経営方針・目標と今後のプラン」末尾に課題＆方針本文
    (10, 3, 0, "{{issues}}"),
    # 補助事業計画: 「2. 販路開拓等の取組内容」末尾
    (11, 1, 0, "{{plan}}"),
    # 補助事業計画: 「4. 補助事業の効果」末尾
    (11, 3, 0, "{{effect}}"),
]


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"source template not found: {SRC}")

    doc = Document(str(SRC))

    for ti, ri, ci, placeholder in INSERTIONS:
        cell = doc.tables[ti].rows[ri].cells[ci]
        cell.add_paragraph(placeholder)
        print(f"injected {placeholder} -> table[{ti}].rows[{ri}].cells[{ci}]")

    doc.save(str(DST))
    print(f"saved: {DST.relative_to(REPO_ROOT)}  ({DST.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
