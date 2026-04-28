# ============================================================
# docx_writer.py
#   Word(.docx) テンプレートのプレースホルダ
#   ({{company}} など) を実値に差し替えて、新しい docx を保存する。
#
#   ファイル名がもとは docx.py だったが、サードパーティライブラリ
#   "python-docx" の import 名 (= docx) と完全に衝突するため
#   docx_writer.py にリネームしてある。
# ============================================================

from pathlib import Path  # 入出力パスの型注釈に使用

# python-docx ライブラリ。Document() で .docx ファイルを開ける。
from docx import Document


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """1 段落 (paragraph) の中でプレースホルダを実値に置換する。

    【背景】
    Word の段落は内部で「run」という小単位に分かれている。
    例: "売上は{{company}}の柱です" が ["売上は", "{{", "company", "}}", "の柱です"]
    のように複数 run に裂けていることがよくある (太字や色替えの境目で分裂)。

    paragraph.text = ... のように代入すると、python-docx は全 run を捨てて
    1 つの run にまとめ直すため、フォント・サイズ・色などの書式が消える。
    そのため可能な限り run 単位で置換し、書式を温存する。
    """

    # まず各 run の中だけで置換を試みる(プレースホルダが 1 run に収まっているケース)
    for run in paragraph.runs:
        for key, value in mapping.items():
            if key in run.text:
                run.text = run.text.replace(key, value)

    # ここまでで置換しきれていない = プレースホルダが run をまたいでいるケース。
    # その場合だけ最後の手段として「最初の run にまとめる」フォールバックを使う。
    # (書式は最初の run のものに揃ってしまうが、未置換のまま残すよりはマシ)
    if any(key in paragraph.text for key in mapping):
        merged = paragraph.text
        for key, value in mapping.items():
            merged = merged.replace(key, value)

        if paragraph.runs:
            paragraph.runs[0].text = merged
            # 2 番目以降の run はテキストを空にして二重表示を防ぐ
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = merged


def fill_docx(template_path: str | Path, output_path: str | Path, sections: dict[str, str]) -> None:
    """テンプレ docx を読み込み、プレースホルダを差し替えて output_path に保存する。

    Args:
        template_path: 入力テンプレートの .docx (プレースホルダ入り)
        output_path:   出力先 .docx
        sections:      {"company": ..., "issues": ..., "plan": ..., "effect": ...}
    """

    # Document は str しか受け付けないバージョンがあるので念のため str() 化
    doc = Document(str(template_path))

    # テンプレ側のプレースホルダ表記 → 差し込みたい値 のマップ
    mapping = {
        "{{company}}": sections["company"],
        "{{issues}}": sections["issues"],
        "{{plan}}": sections["plan"],
        "{{effect}}": sections["effect"],
    }

    # ① 本文の段落を走査
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, mapping)

    # ② 表(table)の中の段落も走査。docx は本文と表で paragraph 列が分かれている。
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, mapping)

    # 保存。output_path も str に揃える。
    doc.save(str(output_path))
